import os
import re

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Inches, Pt  # type: ignore


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def parse_markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    doc = Document()

    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style defaults
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_center = False
    in_table = False
    table_headers = []
    table_rows = []
    in_blockquote = False
    blockquote_text = []

    for line in lines:
        stripped = line.strip()

        # Handle center div tag
        if '<div align="center">' in stripped:
            in_center = True
            continue
        if "</div>" in stripped:
            in_center = False
            continue

        # Handle tables
        if stripped.startswith("|"):
            # It's a table row
            # Split cells
            cells = [c.strip() for c in stripped.split("|")[1:-1]]

            # Check if it's separator row (e.g. |:---|:---|)
            if all(re.match(r"^:?-+:?$", c) for c in cells):
                continue

            if not in_table:
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            continue
        else:
            if in_table:
                # End of table, write it to docx
                if table_headers:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = "Light Shading Accent 1"

                    # Set header row
                    hdr_cells = table.rows[0].cells
                    for i, text in enumerate(table_headers):
                        hdr_cells[i].text = text
                        set_cell_shading(hdr_cells[i], "1F4E79")
                        set_cell_margins(hdr_cells[i])
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = None  # default white or light

                    for row_data in table_rows:
                        # Pad row_data if not matching column length
                        while len(row_data) < len(table_headers):
                            row_data.append("")
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row_data[: len(table_headers)]):
                            row_cells[i].text = val
                            set_cell_margins(row_cells[i])
                            # Check for specific formatting like ✅ or 🔄
                            for paragraph in row_cells[i].paragraphs:
                                # Simple formatting
                                for run in paragraph.runs:
                                    if "✅" in run.text:
                                        run.font.bold = True
                doc.add_paragraph()  # empty space after table
                in_table = False
                table_headers = []
                table_rows = []

        # Handle blockquotes
        if stripped.startswith(">"):
            in_blockquote = True
            bq_content = stripped[1:].strip()
            # Remove alert markup like [!NOTE], [!IMPORTANT]
            bq_content = re.sub(r"^\[!(NOTE|IMPORTANT|WARNING|TIP|CAUTION)\]\s*", "", bq_content)
            blockquote_text.append(bq_content)
            continue
        else:
            if in_blockquote:
                # Write blockquote
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(" ".join(blockquote_text))
                run.italic = True
                run.font.color.rgb = None
                in_blockquote = False
                blockquote_text = []

        # Skip horizontal rules
        if stripped == "---":
            p = doc.add_paragraph()
            p.add_run("____________________________________________________").font.color.rgb = None
            continue

        # Skip empty lines
        if not stripped:
            continue

        # Handle headers
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            h = doc.add_heading(text, level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            continue
        elif stripped.startswith("## "):
            text = stripped[3:].strip()
            h = doc.add_heading(text, level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            continue
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            h = doc.add_heading(text, level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)
            continue

        # Handle bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            # Format bold text e.g. **Text**
            parts = re.split(r"\*\*(.*?)\*\*", text)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.font.bold = True
            continue

        # Handle regular paragraphs
        p = doc.add_paragraph()
        if in_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse inline markdown (bolding only for simplicity)
        parts = re.split(r"\*\*(.*?)\*\*", stripped)
        for idx, part in enumerate(parts):
            # Parse links out for clean text
            part_clean = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", part)
            # Parse latex out
            part_clean = part_clean.replace("$", "")
            run = p.add_run(part_clean)
            if idx % 2 == 1:
                run.font.bold = True

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")


if __name__ == "__main__":
    parse_markdown_to_docx(
        "/home/binhchuoiz/Projects/Thesis/AI_Security_Graph/docs/SENTINEL_Progress_Report.md",
        "/home/binhchuoiz/Projects/Thesis/AI_Security_Graph/docs/SENTINEL_Progress_Report.docx",
    )
